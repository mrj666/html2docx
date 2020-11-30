from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()
# document.add_heading('xxx', 0)
table = document.add_table(rows=6, cols=3, style='Table Grid')
table.cell(0, 0).width = Cm(2)
table.cell(0, 1).width = Cm(6)

for row in table.rows:
    # row.cells[0].text = 'hh'
    row.height = Cm(2)
    cell_new = row.cells[0].paragraphs[0]
    cell_new.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_new.add_run()
    run.add_picture('111.png', height=Inches(0.5))

    cell_new = row.cells[1].paragraphs[0]
    cell_new.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_new.add_run()
    run.add_picture('666.png', height=Inches(0.5))


document.save('demo.docx')
